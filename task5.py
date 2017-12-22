from docx import Document, text
from docx.enum.text import WD_ALIGN_PARAGRAPH

document = Document()
p = document.add_paragraph('[Logo of Affiliation]')
p.add_run(' '*85+'RedClock, PhD')

##

document.add_heading('Petition Letter', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER

p = document.add_paragraph()
run = p.add_run('Summary: ')
run.font.underline = True
run.bold = True
p.add_run('My Curriculum Vitae is in ')
p.add_run('[Exhibit O1]. ').bold = True
p.add_run('I am a scientist in ')
p.add_run('XXXX').font.underline = True
p.add_run('field. ')
p.add_run('I obtained my Ph.D. degree from XXXX in yyyy ')
p.add_run('(Ph.D. diploma in ')
p.add_run('[Exhibit O2]. ').bold = True
p.add_run('I am now a XXXX at XXXX (job verification and employer information in ')
p.add_run('[Exhibit Q3]. ').bold = True
p.add_run('My legal documents are in ')
p.add_run('[Exhibit O4, O5].')

p = document.add_paragraph()
run = p.add_run('Expert Testimony letters')
run.bold = True
run.font.underline = True
p.add_run(': Please see ')
p.add_run('[Exhibit O6-O11].').bold = True

document.add_heading('1. I am an alien of extraordinary ability', level=1).bold = True
h = document.add_heading('', level=1)
run = h.add_run('1-A. Authorship of Scholarly Articles')
run.bold = True
run.font.underline = True
h.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = document.add_paragraph('As my CV in [Exhibit O1] shows, and summarized below: ')
p.add_run('I have exceptional publication records in multiple prestigious peer-reviewed journals with worldwide circulation:')

##

p = document.add_paragraph('A-a. My publications are extraordinary, as they ')
p.add_run('appeared in the leading journals and proceedings').font.underline = True
p.add_run(' with ')
p.add_run('international circulations.').font.underline = True

p = document.add_paragraph('Publication in scientific world is peer-reviewed and highly competitive. ')
p.add_run('Therefore, only the ')
p.add_run('most innovative').bold = True
p.add_run(' and ')
p.add_run('most rigorous').bold = True
p.add_run(' works get accepted and published in top-ranking international journals.')
p.add_run(' As TABLE 1 shows below, ')
p.add_run('my publications appeared in international journals that are very top in related areas.').bold = True

recordset = [{'Jnl': 'Journal XYZ', 'IF': '4.xxxx', 'Rk': 'Top 0.25% (#3/1193 in the xyz field)', 'Ex': '[Exhibit 33]'},
{'Jnl': 'Journal XZY', 'IF': '4.xxxx', 'Rk': 'Top 0.75% (#3/395 in the xzy field)', 'Ex': '[Exhibit 36]'},
{'Jnl': 'Journal YXZ', 'IF': '3.xxxx', 'Rk': 'Top 1.00% (#3/276 in the yxz field)', 'Ex': '[Exhibit 39]'},
{'Jnl': 'Journal YZX', 'IF': '3.xxxx', 'Rk': 'Top 3.00% (#10/354 in the yzx field)', 'Ex': '[Exhibit 42]'},
{'Jnl': 'Journal ZXY', 'IF': '2.xxxx', 'Rk': 'Top 15.0% (#38/276 in the zxy field)', 'Ex': '[Exhibit 44]'}]

document.add_paragraph('TABLE 1. Impact Factors (IFs) and ranks of journals where I published scholarly articles.')
table = document.add_table(rows=1, cols=4, style='Light Shading Accent 3')
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Journals that Published my Articles'
hdr_cells[1].text = 'Impact Factor'
hdr_cells[2].text = 'Journal Ranking'
hdr_cells[3].text = 'Exhibit'
for item in recordset:
    row_cells = table.add_row().cells
    row_cells[0].text = item['Jnl']
    row_cells[1].text = item['IF']
    row_cells[2].text = item['Rk']
    row_cells[3].text = item['Ex']

p = document.add_paragraph('== Below are more detailed introductions of these journals and proceedings ==')
p.add_run('Journal XYZ').font.underline = True
p.add_run(' ranks ')
p.add_run('top 0.25% ').bold = True
p.add_run('(#3 from 1193) journals in the xyz field [')
p.add_run('Exhibit A33').bold = True
p.add_run(']. "The journal publishes ')
p.add_run('original contributions ').bold = True
p.add_run('on xxx." [')
p.add_run('Exhibit A34').bold = True
p.add_run(']. It is published by IEEE, ')
p.add_run('"the world\'s largest technical professional society", ').bold = True
p.add_run('with ')
p.add_run('"over 395,000 members in 160 countries". [').bold = True
p.add_run('Exhibit A35').bold = True
p.add_run('].')

p = document.add_paragraph('--Petition Letter, Page 1 of 19--')
p.bold = True
p.alignment = WD_ALIGN_PARAGRAPH.CENTER


document.save('task5.docx')

